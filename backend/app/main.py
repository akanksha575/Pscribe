import os
import uuid
import shutil
from supabase import create_client
from fastapi import FastAPI, Form, Depends, HTTPException, UploadFile, File
from dotenv import load_dotenv
import logging
import os

env_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), ".env")
load_dotenv(dotenv_path=env_path)

supabase_url = os.getenv("SUPABASE_URL")
supabase_anon_key = os.getenv("SUPABASE_ANON_KEY")
supabase_service_role_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
if not supabase_url or not supabase_anon_key:
    print("WARNING: Supabase URL or Key is missing. Check your .env file!")
supabase = create_client(supabase_url or "", supabase_anon_key or "")

# Service-role client for admin operations (signup user creation)
if supabase_service_role_key:
    supabase_admin = create_client(supabase_url or "", supabase_service_role_key)
else:
    supabase_admin = supabase
    print("WARNING: SUPABASE_SERVICE_ROLE_KEY missing. Signup may not auto-confirm users.")




from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.background import BackgroundTask
from typing import Optional, List
from datetime import date as py_date
from datetime import datetime as dt_datetime 

from sqlalchemy.orm import Session, joinedload, selectinload
from sqlalchemy import or_ 

from . import db_models, schemas # schemas.PatientInfoForNote will be used
from .db_models import create_db_tables, get_db, Patient, MedicalDocument, Encounter

from .scribe import PCScribe, AVAILABLE_TEMPLATES, DEFAULT_TEMPLATE_ID
from .core.config import OPENAI_API_KEY, PATIENT_DATA_DIR

from docx import Document
import datetime
import tempfile
# import speech_recognition as sr
import openai

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="PCScribe API", version="1.6.0") # Updated version
app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

pcscribe_instance = PCScribe(api_key=OPENAI_API_KEY) # Pass OpenAI key
@app.post("/api/auth/login")
async def login(username: str = Form(...), password: str = Form(...)):
    try:
        # Use the correct method for password authentication in supabase-py v2
        response = supabase.auth.sign_in_with_password({"email": username, "password": password})
        # response contains a 'user' key when successful
        if getattr(response, 'user', None) or (isinstance(response, dict) and response.get('user')):
            return {"status": "ok"}
        else:
            raise HTTPException(status_code=401, detail="Invalid credentials")
    except Exception as e:
        error_msg = str(e)
        if "Invalid login credentials" in error_msg or "Invalid credentials" in error_msg:
            raise HTTPException(status_code=401, detail="Invalid login credentials")
        raise HTTPException(status_code=500, detail=error_msg)

@app.post("/api/auth/signup")
async def signup(email: str = Form(...), password: str = Form(...), full_name: str = Form("")):
    try:
        # Use admin client to create user with auto-confirm
        response = supabase_admin.auth.admin.create_user({
            "email": email,
            "password": password,
            "email_confirm": True,
            "user_metadata": {"full_name": full_name}
        })
        if getattr(response, 'user', None):
            return {"status": "ok", "message": "Account created successfully. You can now sign in."}
        else:
            raise HTTPException(status_code=400, detail="Could not create account.")
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Signup error: {error_msg}")
        if "already been registered" in error_msg or "already exists" in error_msg or "duplicate" in error_msg.lower():
            raise HTTPException(status_code=409, detail="An account with this email already exists. Please sign in instead.")
        if "password" in error_msg.lower() and ("short" in error_msg.lower() or "weak" in error_msg.lower() or "length" in error_msg.lower()):
            raise HTTPException(status_code=400, detail="Password must be at least 6 characters long.")
        raise HTTPException(status_code=500, detail=f"Signup failed: {error_msg}")

@app.on_event("startup")
async def startup_event():
    logger.info("Creating database tables if they don't exist...")
    create_db_tables()
    logger.info("Database tables checked/created.")

# --- Patient Endpoints ---
@app.post("/api/patients/", response_model=schemas.PatientDisplay, status_code=201)
async def create_patient(
    patient_data: schemas.PatientCreate,
    db: Session = Depends(get_db)
):
    if not patient_data.first_name or not patient_data.last_name:
        raise HTTPException(status_code=422, detail="First name and last name are required.")

    if patient_data.account_no:
        existing_patient_by_acc = db.query(db_models.Patient).filter(db_models.Patient.account_no == patient_data.account_no).first()
        if existing_patient_by_acc:
            raise HTTPException(status_code=409, detail=f"Patient with account number {patient_data.account_no} already exists.")

    db_patient_dict = patient_data.dict(exclude_unset=True)
    db_patient = db_models.Patient(**db_patient_dict)

    db.add(db_patient)
    try:
        db.commit()
        db.refresh(db_patient)
    except Exception as e:
        db.rollback()
        logger.error(f"Error creating patient: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Could not create patient in database: {str(e)}")
    return db_patient

@app.get("/api/patients/", response_model=List[schemas.PatientDisplay])
async def list_patients(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    patients = db.query(db_models.Patient).order_by(db_models.Patient.last_name, db_models.Patient.first_name).offset(skip).limit(limit).all()
    return patients

@app.get("/api/patients/{patient_id}", response_model=schemas.PatientDisplay)
async def get_patient(patient_id: int, db: Session = Depends(get_db)):
    db_patient = db.query(db_models.Patient).options(selectinload(db_models.Patient.documents)).filter(db_models.Patient.id == patient_id).first()
    if db_patient is None:
        raise HTTPException(status_code=404, detail="Patient not found")
    return db_patient

@app.get("/api/patients/search/", response_model=List[schemas.PatientDisplay]) 
async def search_patients(
    name: Optional[str] = None,
    dob: Optional[str] = None, 
    account_no: Optional[str] = None,
    db: Session = Depends(get_db)
):
    query = db.query(db_models.Patient)
    
    if name:
        name_parts = name.split(',')
        if len(name_parts) > 1:
            last_name_query = name_parts[0].strip()
            first_name_query = name_parts[1].strip()
            query = query.filter(
                db_models.Patient.last_name.ilike(f"%{last_name_query}%"),
                db_models.Patient.first_name.ilike(f"%{first_name_query}%")
            )
        elif name.strip(): 
            name_query = name.strip()
            query = query.filter(
                or_(
                    db_models.Patient.last_name.ilike(f"%{name_query}%"),
                    db_models.Patient.first_name.ilike(f"%{name_query}%")
                )
            )
    
    if dob:
        try:
            dob_date = dt_datetime.strptime(dob, "%Y-%m-%d").date()
        except ValueError:
            try:
                dob_date = dt_datetime.strptime(dob, "%m-%d-%Y").date()
            except ValueError:
                raise HTTPException(status_code=400, detail="Invalid DOB format. Use YYYY-MM-DD or MM-DD-YYYY.")
        query = query.filter(db_models.Patient.dob == dob_date)

    if account_no:
        query = query.filter(db_models.Patient.account_no.ilike(f"%{account_no}%"))
    
    if not name and not dob and not account_no:
        raise HTTPException(status_code=400, detail="Please provide at least one search criteria (name, DOB, or account number).")

    patients = query.order_by(db_models.Patient.last_name, db_models.Patient.first_name).limit(20).all()
    return patients


# --- Document Upload/Download Endpoints ---
@app.post("/api/patients/{patient_id}/documents/", response_model=schemas.MedicalDocumentDisplay)
async def upload_patient_document(
    patient_id: int,
    file: UploadFile = File(...),
    document_type: str = Form(...), 
    db: Session = Depends(get_db)
):
    db_patient = db.query(db_models.Patient).filter(db_models.Patient.id == patient_id).first()
    if not db_patient:
        raise HTTPException(status_code=404, detail="Patient not found to associate document with.")

    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided.")
    
    _, extension = os.path.splitext(file.filename)
    unique_filename_on_disk = f"{uuid.uuid4()}{extension}"
    
    patient_file_subdir = os.path.join(PATIENT_DATA_DIR, str(patient_id), document_type.lower().replace(" ", "_") + "s")
    os.makedirs(patient_file_subdir, exist_ok=True)
    
    file_path_on_disk = os.path.join(patient_file_subdir, unique_filename_on_disk)
    
    try:
        with open(file_path_on_disk, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        logger.error(f"Could not save uploaded file {file.filename} for patient {patient_id}: {e}")
        raise HTTPException(status_code=500, detail="Could not save uploaded file.")
    finally:
        await file.close() 

    db_relative_file_path = os.path.join(str(patient_id), document_type.lower().replace(" ", "_") + "s", unique_filename_on_disk)

    db_document = db_models.MedicalDocument(
        patient_id=patient_id,
        filename=unique_filename_on_disk,
        original_filename=file.filename,
        document_type=document_type,
        file_path=db_relative_file_path,
        content_type=file.content_type or "application/octet-stream"
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    return db_document

@app.get("/api/documents/{document_id}", response_class=FileResponse)
async def download_document(document_id: int, db: Session = Depends(get_db)): 
    db_document = db.query(db_models.MedicalDocument).filter(db_models.MedicalDocument.id == document_id).first()
    if not db_document:
        raise HTTPException(status_code=404, detail="Document not found")

    full_file_path = os.path.join(PATIENT_DATA_DIR, db_document.file_path)
    if not os.path.exists(full_file_path):
        logger.error(f"Document file not found on disk: {full_file_path} for document ID {document_id}")
        raise HTTPException(status_code=404, detail="File not found on server.")

    return FileResponse(
        path=full_file_path,
        filename=db_document.original_filename,
        media_type=db_document.content_type
    )

# --- Encounter Endpoints ---
@app.get("/api/patients/{patient_id}/encounters/", response_model=List[schemas.EncounterDisplay])
async def list_patient_encounters(patient_id: int, db: Session = Depends(get_db)):
    patient = db.query(db_models.Patient).filter(db_models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")

    encounters = db.query(Encounter)\
        .filter(Encounter.patient_id == patient_id)\
        .options(
            joinedload(Encounter.transcript_document),
            joinedload(Encounter.generated_note_document)
        )\
        .order_by(Encounter.appointment_date.desc(), Encounter.created_at.desc())\
        .all()
    
    results = []
    for enc in encounters:
        enc_display = schemas.EncounterDisplay.from_orm(enc)
        enc_display.patient_first_name = patient.first_name
        enc_display.patient_last_name = patient.last_name
        enc_display.patient_dob = patient.dob
        results.append(enc_display)
    return results


# --- Note Generation Endpoint (Updated) ---
@app.post("/api/generate-note/", response_model=schemas.GenerateNoteResponse)
async def generate_note_endpoint(
    patient_id: Optional[int] = Form(None),
    transcript_document_id: Optional[int] = Form(None),
    file: Optional[UploadFile] = File(None), 
    provider_name: str = Form("Dr. Default"),
    template_id: str = Form(DEFAULT_TEMPLATE_ID),
    appointment_date_str: Optional[str] = Form(None),
    patient_first_name: Optional[str] = Form(None),
    patient_last_name: Optional[str] = Form(None),
    patient_dob: Optional[str] = Form(None), 
    patient_sex: Optional[str] = Form(None),
    patient_account_no: Optional[str] = Form(None),
    # Add all other patient fields from schemas.PatientCreate if they are part of the form
    patient_address1: Optional[str] = Form(None),
    patient_city: Optional[str] = Form(None),
    patient_state: Optional[str] = Form(None),
    patient_zip_code: Optional[str] = Form(None), # Match model field name if aliased 'zip' in schema
    patient_cell_phone: Optional[str] = Form(None),
    patient_email: Optional[str] = Form(None),
    patient_race: Optional[str] = Form(None),
    patient_ethnicity: Optional[str] = Form(None),
    patient_marital_status: Optional[str] = Form(None),
    patient_pcp: Optional[str] = Form(None),
    patient_ec_name: Optional[str] = Form(None),
    patient_ec_relation: Optional[str] = Form(None),
    patient_ec_phone: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    db_patient: Optional[Patient] = None
    transcript_text: Optional[str] = None
    source_transcript_document: Optional[MedicalDocument] = None
    encounter: Optional[Encounter] = None
    
    appointment_date_obj: Optional[py_date] = None
    if appointment_date_str:
        try:
            appointment_date_obj = dt_datetime.strptime(appointment_date_str, "%Y-%m-%d").date()
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid appointment_date format. Use YYYY-MM-DD.")

    if patient_id:
        db_patient = db.query(Patient).filter(Patient.id == patient_id).first()
        if not db_patient:
            raise HTTPException(status_code=404, detail=f"Patient with ID {patient_id} not found.")
    elif file: 
        if not patient_first_name or not patient_last_name:
            raise HTTPException(status_code=422, detail="First name and last name are required for a new patient.")
        
        new_patient_data_dict = {
            "first_name": patient_first_name, "last_name": patient_last_name, "dob": patient_dob,
            "sex": patient_sex, "account_no": patient_account_no,
            "address1": patient_address1, "city": patient_city, "state": patient_state,
            "zip_code": patient_zip_code, "cell_phone": patient_cell_phone, "email": patient_email,
            "race": patient_race, "ethnicity": patient_ethnicity, "marital_status": patient_marital_status,
            "pcp": patient_pcp, "ec_name": patient_ec_name, "ec_relation": patient_ec_relation, "ec_phone": patient_ec_phone
        }
        # Filter out None values before passing to Pydantic model for creation
        filtered_new_patient_data = {k: v for k, v in new_patient_data_dict.items() if v is not None}
        new_patient_data = schemas.PatientCreate(**filtered_new_patient_data)
        
        if new_patient_data.account_no:
            existing_by_acc = db.query(Patient).filter(Patient.account_no == new_patient_data.account_no).first()
            if existing_by_acc:
                raise HTTPException(status_code=409, detail=f"A patient with account number {new_patient_data.account_no} already exists. Please search and select the patient, or use a different account number.")

        db_patient = Patient(**new_patient_data.dict(exclude_unset=True)) # Use .dict() from Pydantic model
        db.add(db_patient)
        try:
            db.commit()
            db.refresh(db_patient)
        except Exception as e_pat_create:
            db.rollback()
            logger.error(f"Error creating new patient during note generation: {e_pat_create}")
            raise HTTPException(status_code=500, detail=f"Could not create new patient: {str(e_pat_create)}")

    if transcript_document_id:
        source_transcript_document = db.query(MedicalDocument).filter(
            MedicalDocument.id == transcript_document_id,
            MedicalDocument.document_type == "transcript"
        ).first()
        if not source_transcript_document:
            raise HTTPException(status_code=404, detail=f"Transcript document with ID {transcript_document_id} not found or not a transcript.")

        if db_patient and source_transcript_document.patient_id != db_patient.id:
            raise HTTPException(status_code=400, detail="Transcript document does not belong to the specified patient.")
        elif not db_patient:
            db_patient = source_transcript_document.patient 

        transcript_file_full_path = os.path.join(PATIENT_DATA_DIR, source_transcript_document.file_path)
        if not os.path.exists(transcript_file_full_path):
            logger.error(f"Transcript file missing: {transcript_file_full_path}")
            raise HTTPException(status_code=500, detail="Transcript file is missing from storage.")
        
        try:
            transcript_text = pcscribe_instance.extract_text_from_docx(transcript_file_full_path)
        except Exception as e_extract:
            logger.error(f"Error extracting text from stored transcript {transcript_document_id}: {e_extract}")
            raise HTTPException(status_code=500, detail=f"Could not extract text from stored transcript: {str(e_extract)}")

        encounter = db.query(Encounter).filter(Encounter.transcript_document_id == source_transcript_document.id).first()
        if not encounter: 
            encounter = Encounter(
                patient_id=db_patient.id,
                transcript_document_id=source_transcript_document.id,
                appointment_date=appointment_date_obj or source_transcript_document.upload_date.date() 
            )
            db.add(encounter)
            db.commit() 
            db.refresh(encounter)

    elif file: 
        if not db_patient:
             raise HTTPException(status_code=400, detail="Patient context (existing ID or new patient data) is required when uploading a new transcript.")
        
        # Debug logging
        logger.info(f"Received file: {file.filename}, content_type: {file.content_type}")
        
        # Handle both .docx files and text content from recorded transcripts
        if file.filename and file.filename.lower().endswith(".docx"):
            logger.info("Processing as .docx file")
            # Process as .docx file
            _, extension = os.path.splitext(file.filename)
            transcript_disk_filename = f"{uuid.uuid4()}{extension}"
            transcript_subdir = os.path.join(PATIENT_DATA_DIR, str(db_patient.id), "transcripts")
            os.makedirs(transcript_subdir, exist_ok=True)
            transcript_disk_path = os.path.join(transcript_subdir, transcript_disk_filename)

            try:
                with open(transcript_disk_path, "wb") as buffer:
                    shutil.copyfileobj(file.file, buffer)
                transcript_text = pcscribe_instance.extract_text_from_docx(transcript_disk_path)
                logger.info(f"Extracted text from .docx: {len(transcript_text)} characters")
            except Exception as e_upload_extract:
                logger.error(f"Error saving or extracting text from new transcript: {e_upload_extract}")
                if os.path.exists(transcript_disk_path): os.remove(transcript_disk_path)
                raise HTTPException(status_code=500, detail=f"Could not process uploaded transcript: {str(e_upload_extract)}")
            finally:
                await file.close()

            source_transcript_document = MedicalDocument(
                patient_id=db_patient.id,
                filename=transcript_disk_filename,
                original_filename=file.filename,
                document_type="transcript",
                file_path=os.path.join(str(db_patient.id), "transcripts", transcript_disk_filename),
                content_type=file.content_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif file.filename and file.filename.lower().endswith(".txt"):
            logger.info("Processing as .txt file (recorded transcript)")
            # Process as text content (from recorded transcripts)
            try:
                transcript_text = await file.read()
                transcript_text = transcript_text.decode('utf-8')
                logger.info(f"Read text content: {len(transcript_text)} characters")
                await file.close()
                
                # Save as .docx for consistency
                doc = Document()
                doc.add_heading("Transcript", 0)
                doc.add_paragraph(transcript_text)
                
                transcript_disk_filename = f"{uuid.uuid4()}.docx"
                transcript_subdir = os.path.join(PATIENT_DATA_DIR, str(db_patient.id), "transcripts")
                os.makedirs(transcript_subdir, exist_ok=True)
                transcript_disk_path = os.path.join(transcript_subdir, transcript_disk_filename)
                doc.save(transcript_disk_path)
                logger.info(f"Saved transcript as .docx: {transcript_disk_path}")
                
                source_transcript_document = MedicalDocument(
                    patient_id=db_patient.id,
                    filename=transcript_disk_filename,
                    original_filename=file.filename,
                    document_type="transcript",
                    file_path=os.path.join(str(db_patient.id), "transcripts", transcript_disk_filename),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e_text_process:
                logger.error(f"Error processing text transcript: {e_text_process}")
                raise HTTPException(status_code=500, detail=f"Could not process text transcript: {str(e_text_process)}")
        else:
            logger.error(f"Invalid file type: {file.filename}")
            await file.close()
            raise HTTPException(status_code=400, detail="Invalid transcript file type. Only .docx or .txt accepted.")

        db.add(source_transcript_document)
        db.commit() 
        db.refresh(source_transcript_document)

        encounter = Encounter(
            patient_id=db_patient.id,
            transcript_document_id=source_transcript_document.id,
            appointment_date=appointment_date_obj or py_date.today()
        )
        db.add(encounter)
        db.commit() 
        db.refresh(encounter)
    else:
        raise HTTPException(status_code=400, detail="Either an existing transcript_document_id or a new transcript file must be provided.")

    if not transcript_text:
        raise HTTPException(status_code=500, detail="Failed to load transcript text for processing.")
    if not db_patient: 
        raise HTTPException(status_code=500, detail="Patient context could not be established.")

    patient_info_to_pass = {
        "first_name": db_patient.first_name, "last_name": db_patient.last_name,
        "dob": db_patient.dob.strftime("%Y-%m-%d") if db_patient.dob else None, # Scribe expects YYYY-MM-DD
        "sex": db_patient.sex, "account_no": db_patient.account_no,
        "address1": db_patient.address1, "city": db_patient.city, "state": db_patient.state,
        "zip": db_patient.zip_code, "cell_phone": db_patient.cell_phone, "email": db_patient.email,
        "race": db_patient.race, "ethnicity": db_patient.ethnicity,
        "marital_status": db_patient.marital_status, "pcp": db_patient.pcp,
        "ec_name": db_patient.ec_name, "ec_relation": db_patient.ec_relation, "ec_phone": db_patient.ec_phone
    }
    # Filter None values for set_patient_info, as scribe.py HTML templates handle N/A
    pcscribe_instance.set_patient_info({k:v for k,v in patient_info_to_pass.items() if v is not None})
    pcscribe_instance.set_template(template_id)

    try:
        html_note_content, generated_docx_bytes, transient_note_id = pcscribe_instance.process_transcript_text(
            transcript_text, provider_name, anonymize_pii=True 
        )
    except Exception as e_scribe:
        logger.error(f"Error during PCScribe processing: {e_scribe}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Note generation failed: {str(e_scribe)}")

    active_template_name = AVAILABLE_TEMPLATES.get(pcscribe_instance.active_template_id, {}).get("name", "Note").replace(' ','_')
    generated_note_orig_filename = f"{db_patient.last_name}_{db_patient.first_name}_Encounter_{encounter.id}_{active_template_name}_{transient_note_id[:8]}.docx"
    generated_note_disk_filename = f"{uuid.uuid4()}.docx" 
    
    generated_notes_subdir = os.path.join(PATIENT_DATA_DIR, str(db_patient.id), "generated_notes")
    os.makedirs(generated_notes_subdir, exist_ok=True)
    generated_note_disk_path = os.path.join(generated_notes_subdir, generated_note_disk_filename)

    try:
        with open(generated_note_disk_path, "wb") as f:
            f.write(generated_docx_bytes)
    except Exception as e_save_note:
        logger.error(f"Could not save generated note DOCX to disk: {e_save_note}")
        raise HTTPException(status_code=500, detail="Failed to save generated note file.")

    db_generated_note_document = MedicalDocument(
        patient_id=db_patient.id,
        filename=generated_note_disk_filename,
        original_filename=generated_note_orig_filename,
        document_type="generated_note",
        file_path=os.path.join(str(db_patient.id), "generated_notes", generated_note_disk_filename),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    db.add(db_generated_note_document)
    
    encounter.generated_note_document = db_generated_note_document 
    if appointment_date_obj : encounter.appointment_date = appointment_date_obj 

    # Construct PatientInfoForNote for the response
    patient_details_for_response = schemas.PatientInfoForNote(
        account_no=db_patient.account_no or "N/A",
        name=f"{db_patient.last_name or ''}, {db_patient.first_name or ''}".strip(", "),
        dob=db_patient.dob.strftime("%Y-%m-%d") if db_patient.dob else "N/A",
        sex=db_patient.sex or "N/A",
        address= (f"{db_patient.address1 or ''} {db_patient.city or ''} {db_patient.state or ''} {db_patient.zip_code or ''}").replace("  ", " ").strip() or "N/A",
        cell_phone=db_patient.cell_phone or "N/A",
        email=db_patient.email or "N/A"
    )
    # Ensure address isn't just spaces if all parts are empty
    if not patient_details_for_response.address.replace("N/A","").strip():
        patient_details_for_response.address = "N/A"


    try:
        db.commit()
        db.refresh(db_generated_note_document)
        db.refresh(encounter)
        db.refresh(db_patient) # Refresh patient in case it was newly created and we need its final state
    except Exception as e_final_commit:
        db.rollback()
        logger.error(f"Error committing generated note and encounter update: {e_final_commit}")
        if os.path.exists(generated_note_disk_path): os.remove(generated_note_disk_path)
        raise HTTPException(status_code=500, detail="Failed to finalize saving generated note to database.")
        
    return schemas.GenerateNoteResponse(
        message=f"Clinical note generated successfully for patient {db_patient.first_name} {db_patient.last_name} (Encounter ID: {encounter.id}).",
        html_preview=html_note_content,
        template_used=AVAILABLE_TEMPLATES.get(pcscribe_instance.active_template_id, {}).get("name", "Selected Template"),
        patient_id=db_patient.id,
        encounter_id=encounter.id,
        generated_document_id=db_generated_note_document.id,
        patient_details=patient_details_for_response # Include patient details
    )


# --- Health and Templates Endpoints ---
@app.get("/api/health", response_model=schemas.HealthCheckResponse)
async def health_check():
    return {"status": "healthy", "api_key_present": bool(OPENAI_API_KEY)}
@app.get("/api/templates", response_model=schemas.TemplatesResponse)
async def get_templates():
    templates_dict = {template_id: data["name"] for template_id, data in AVAILABLE_TEMPLATES.items()}
    return schemas.TemplatesResponse(templates=templates_dict)

@app.post("/api/transcribe-audio/")
async def transcribe_audio(
    audio: UploadFile = File(...),
    patient_first_name: str = Form(...),
    patient_last_name: str = Form(...),
    patient_id: int = Form(...),
    date: str = Form(...)
):
    # Save uploaded audio to a temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as temp_audio:
        temp_audio.write(await audio.read())
        temp_audio_path = temp_audio.name

    transcript_text = ""
    try:
        with open(temp_audio_path, "rb") as audio_file:
            transcript_response = openai.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )
            transcript_text = transcript_response.text
    except Exception as e:
        os.remove(temp_audio_path)
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")
    os.remove(temp_audio_path)

    # Create .docx file
    doc = Document()
    doc.add_heading(f"Transcript for {patient_last_name}, {patient_first_name}", 0)
    doc.add_paragraph(transcript_text)
    filename = f"{patient_last_name}_{patient_first_name}_{date}.docx"
    safe_filename = filename.replace(' ', '_').replace(',', '')
    transcripts_dir = os.path.join(os.path.dirname(__file__), '../../transcripts')
    os.makedirs(transcripts_dir, exist_ok=True)
    docx_path = os.path.join(transcripts_dir, safe_filename)
    doc.save(docx_path)

    # Return JSON with transcript and filename
    return {"transcript": transcript_text, "filename": safe_filename}

@app.get("/api/documents/by-filename/{filename}")
async def download_document_by_filename(filename: str):
    transcripts_dir = os.path.join(os.path.dirname(__file__), '../../transcripts')
    file_path = os.path.join(transcripts_dir, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True, app_dir=os.path.dirname(__file__) or ".")